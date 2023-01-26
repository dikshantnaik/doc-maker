import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


def change_header(file, name, class_name, roll_no):
    # Open the .docx file
    doc = docx.Document(file)

    # Get the header of the first section
    header = doc.sections[0].header

    # Clear the existing header
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            run.clear()
    HEADER_TEXT = f"Name: {name}\n RollNo : {roll_no}\nClass: {class_name}"
    # Add new text to the header
    paragraph = header.add_paragraph(HEADER_TEXT)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        run.font.size = 18
        run.font.bold = True
    # Change the style of the header to "Header 6"
    # paragraph.style = 'Header 6'
    # Save the modified file as a new .docx file
    doc.save("modified_file.docx")
    return "Header successfully changed, aligned to the Right, font size increased, bold font set and style set to 'Header 6'!"


if __name__ == '__main__':
    file = 'unCompa_header.docx'
    name = 'Dikshanntttt'
    class_name = 'Class 10'
    roll_no = '10'
    change_header(file, name, class_name, roll_no)
