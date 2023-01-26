from flask import Flask, request
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import send_file
app = Flask(__name__)


@app.route('/change_header', methods=['POST'])
def change_header():
    # Get the uploaded file and new header text from the POST request
    file = request.files['file']
    name = request.form['name']
    roll_no = request.form['roll_no']
    class_name = request.form['Class']

    doc = Document(file)

    # Get the header of the first section
    header = doc.sections[0].header
    header.is_linked_to_previous = True
    # Clear the existing header
    # for paragraph in header.paragraphs:
    #     paragraph.clear()
    #     for run in paragraph.runs:
    #         run.clear()
    HEADER_TEXT = f"Name: {name}\n RollNo : {roll_no}\nClass: {class_name}"
    # Add new text to the header
    paragraph = header.add_paragraph(HEADER_TEXT)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.style.font.name = "Arial"
    for run in paragraph.runs:
        run.font.size = 18
        # run.font.bold = True
    # Change the style of the header to "Header 6"
    # paragraph.style = 'Header 6'
    # Save the modified file as a new .docx file
    doc.save("modified_file.docx")

    return send_file("../modified_file.docx")
    # return "Header successfully changed!"


# if __name__ == '__main__':
#     app.run(debug=True)
